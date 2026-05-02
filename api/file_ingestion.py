"""
File ingestion polyvalent pour Agent 3 — PDF / Word / Excel / CSV / TXT / JSON.

Objectifs :
  - Extraction fiable (libs natives Python, pas de subprocess)
  - Texte structure preservant la mise en page (tableaux, sections)
  - Hook auto-RAG : si > seuil chars, ingest dans `agent3_embeddings`
  - Fallback propre si une lib optionnelle est absente

Libs optionnelles :
  - PyMuPDF (`fitz`) pour PDF (fallback : pypdf si present)
  - python-docx pour Word .docx
  - openpyxl pour Excel .xlsx (fallback : pandas si dispo)
  - pandas pour meilleure gestion CSV/Excel

Point d'entree : `extract_and_ingest(db, user_id, filepath, filetype, *, auto_rag=True)`.
Retourne : {"text": str, "chars": int, "ingested": bool, "chunks": int, "backend": str}.

Utilise par /api/agent3/upload et par le tool `FILE_ANALYZE` du dispatcher.
"""

from __future__ import annotations

import csv
import json
import logging
import mimetypes
import os
from pathlib import Path
from typing import Any

logger = logging.getLogger("sylea.file_ingestion")


# Seuils
_MAX_TEXT_CHARS = 200_000        # hard cap pour eviter l'explosion memoire
_AUTO_RAG_MIN_CHARS = 2_000       # en dessous, pas la peine de RAG (reste inline)
_EXCEL_MAX_ROWS = 500
_CSV_MAX_ROWS = 500


# ─────────────────────────────────────────────────────────────────────────────
# Detection du type (mime + extension)
# ─────────────────────────────────────────────────────────────────────────────

def detect_mime(filepath: str, provided: str | None = None) -> str:
    """Determine le mime type, preferant le parametre fourni si non vide.

    Note : on override les mappings Windows-specifiques douteux (CSV -> Excel)
    pour avoir des mimes stables cross-platform.
    """
    if provided and "/" in provided:
        return provided
    ext = Path(filepath).suffix.lower()
    # Mapping prioritaire (stable cross-platform, ignore mimetypes OS)
    stable_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".json": "application/json",
    }
    if ext in stable_map:
        return stable_map[ext]
    guessed, _ = mimetypes.guess_type(filepath)
    if guessed:
        return guessed
    return "application/octet-stream"


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> tuple[str, str]:
    """Retourne (texte, backend_utilise). Essaye PyMuPDF puis pypdf."""
    # PyMuPDF (fitz)
    try:
        import fitz  # type: ignore
        doc = fitz.open(str(path))
        parts: list[str] = []
        for page_idx, page in enumerate(doc):
            try:
                text = page.get_text("text")
                if text.strip():
                    parts.append(f"--- Page {page_idx + 1} ---\n{text}")
            except Exception:
                continue
        doc.close()
        return "\n\n".join(parts), "pymupdf"
    except Exception as e:
        logger.debug(f"PyMuPDF indisponible : {e}")

    # Fallback pypdf
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"--- Page {i + 1} ---\n{text}")
            except Exception:
                continue
        return "\n\n".join(parts), "pypdf"
    except Exception as e:
        logger.debug(f"pypdf indisponible : {e}")

    return f"[PDF non extractible : {path.name}, taille {path.stat().st_size} o]", "metadata_only"


# ─────────────────────────────────────────────────────────────────────────────
# Word .docx
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> tuple[str, str]:
    """Retourne (texte, backend). Utilise python-docx."""
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        parts: list[str] = []
        # Paragraphes
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Preserve le niveau de style (titre vs normal)
            style = (para.style.name or "").lower() if para.style else ""
            if "heading" in style or "title" in style:
                parts.append(f"\n## {text}")
            else:
                parts.append(text)
        # Tables
        for t_idx, table in enumerate(doc.tables):
            parts.append(f"\n### Tableau {t_idx + 1}")
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts), "python-docx"
    except ImportError:
        logger.debug("python-docx non installe")
        return f"[Word non extractible (python-docx absent) : {path.name}]", "metadata_only"
    except Exception as e:
        logger.warning(f"docx extraction failed : {e}")
        return f"[Erreur extraction Word : {type(e).__name__}]", "error"


# ─────────────────────────────────────────────────────────────────────────────
# Excel
# ─────────────────────────────────────────────────────────────────────────────

def _extract_xlsx(path: Path) -> tuple[str, str]:
    """Retourne (texte, backend). openpyxl prioritaire, pandas en fallback."""
    # openpyxl : pas de dep pandas
    try:
        from openpyxl import load_workbook  # type: ignore
        wb = load_workbook(str(path), data_only=True, read_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n## Feuille : {sheet_name}")
            rows_collected = 0
            for row in ws.iter_rows(values_only=True):
                if rows_collected >= _EXCEL_MAX_ROWS:
                    parts.append(f"... (tronque a {_EXCEL_MAX_ROWS} lignes)")
                    break
                cells = [("" if v is None else str(v)) for v in row]
                if not any(c.strip() for c in cells):
                    continue
                parts.append(" | ".join(cells))
                rows_collected += 1
        wb.close()
        return "\n".join(parts), "openpyxl"
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"openpyxl extraction failed : {e}")

    # Fallback pandas
    try:
        import pandas as pd  # type: ignore
        xls = pd.ExcelFile(str(path))
        parts = []
        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet, nrows=_EXCEL_MAX_ROWS)
            parts.append(f"\n## Feuille : {sheet}")
            parts.append(df.to_string(index=False))
        return "\n".join(parts), "pandas"
    except ImportError:
        logger.debug("Ni openpyxl ni pandas disponibles pour Excel")
    except Exception as e:
        logger.warning(f"pandas Excel extraction failed : {e}")

    return f"[Excel non extractible : {path.name}]", "metadata_only"


# ─────────────────────────────────────────────────────────────────────────────
# CSV
# ─────────────────────────────────────────────────────────────────────────────

def _extract_csv(path: Path) -> tuple[str, str]:
    """Retourne (texte, backend). Detection delimiter + max rows."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            rows = []
            for i, row in enumerate(reader):
                if i >= _CSV_MAX_ROWS:
                    rows.append(f"... (tronque a {_CSV_MAX_ROWS} lignes)")
                    break
                rows.append(" | ".join(cell.replace("\n", " ") for cell in row))
        return "\n".join(rows), "csv"
    except Exception as e:
        logger.warning(f"csv extraction failed : {e}")
        return f"[CSV non extractible : {path.name}]", "error"


# ─────────────────────────────────────────────────────────────────────────────
# Texte brut
# ─────────────────────────────────────────────────────────────────────────────

def _extract_text(path: Path) -> tuple[str, str]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return text, "text"
    except Exception as e:
        return f"[Fichier texte illisible : {type(e).__name__}]", "error"


def _extract_json(path: Path) -> tuple[str, str]:
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        # Try to pretty-print pour lisibilite LLM
        try:
            parsed = json.loads(raw)
            return json.dumps(parsed, indent=2, ensure_ascii=False), "json"
        except Exception:
            return raw, "text"
    except Exception as e:
        return f"[JSON illisible : {type(e).__name__}]", "error"


# ─────────────────────────────────────────────────────────────────────────────
# Point d'entree principal
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(filepath: str, filetype: str | None = None) -> dict[str, Any]:
    """Extrait le texte d'un fichier.

    Retourne : {"text": str, "chars": int, "backend": str, "mime": str, "truncated": bool}.
    """
    path = Path(filepath)
    if not path.exists() or not path.is_file():
        return {"text": "", "chars": 0, "backend": "missing", "mime": "", "truncated": False}

    mime = detect_mime(filepath, filetype)
    text = ""
    backend = "unknown"

    try:
        if mime == "application/pdf" or path.suffix.lower() == ".pdf":
            text, backend = _extract_pdf(path)
        elif path.suffix.lower() == ".docx" or "wordprocessingml" in mime:
            text, backend = _extract_docx(path)
        elif path.suffix.lower() in (".xlsx", ".xlsm") or "spreadsheetml" in mime:
            text, backend = _extract_xlsx(path)
        elif mime == "text/csv" or path.suffix.lower() == ".csv":
            text, backend = _extract_csv(path)
        elif mime == "application/json" or path.suffix.lower() == ".json":
            text, backend = _extract_json(path)
        elif mime.startswith("text/") or path.suffix.lower() in (".txt", ".md", ".log"):
            text, backend = _extract_text(path)
        elif mime.startswith("image/"):
            # Les images ne sont PAS extraites ici — voir vision_analyze.
            text, backend = (
                f"[Image : {path.name}, {mime}, {path.stat().st_size} o — utilise VISION_ANALYZE]",
                "skip_image",
            )
        else:
            # Dernier recours : lecture brute
            text, backend = _extract_text(path)
    except Exception as e:
        logger.exception(f"extract_text crash on {path.name}: {e}")
        text = f"[Erreur extraction {path.name} : {type(e).__name__}]"
        backend = "error"

    truncated = len(text) > _MAX_TEXT_CHARS
    if truncated:
        text = text[:_MAX_TEXT_CHARS] + f"\n\n[... tronque a {_MAX_TEXT_CHARS} chars]"

    return {
        "text": text,
        "chars": len(text),
        "backend": backend,
        "mime": mime,
        "truncated": truncated,
    }


def extract_only(filepath: str, filetype: str | None = None) -> str:
    """Helper raccourci : retourne juste le texte extrait (str).

    Utilise par les routes qui veulent juste le contenu sans metadata.
    Equivalent a extract_text(...)["text"] mais avec API plus simple.
    """
    return extract_text(filepath, filetype).get("text", "")


async def extract_and_ingest(
    db: Any,
    user_id: str | None,
    filepath: str,
    filetype: str | None = None,
    *,
    auto_rag: bool = True,
    source_ref: str | None = None,
) -> dict[str, Any]:
    """Extrait + optionnellement ingest dans le RAG.

    Retourne : {text, chars, ingested, chunks, backend, mime, truncated}.
    """
    result = extract_text(filepath, filetype)
    text = result["text"]

    ingested = False
    chunks = 0

    if auto_rag and user_id and len(text) >= _AUTO_RAG_MIN_CHARS:
        try:
            from api.rag import ingest_text
            ref = source_ref or os.path.basename(filepath)
            rag_result = await ingest_text(
                db, user_id, text,
                source_type="upload",
                source_ref=ref,
                scope="global",
                metadata={"backend": result["backend"], "mime": result["mime"]},
            )
            chunks = rag_result.get("ingested_chunks", 0)
            ingested = chunks > 0
        except Exception as e:
            logger.warning(f"auto-RAG ingest failed : {e}")

    result["ingested"] = ingested
    result["chunks"] = chunks
    return result


__all__ = [
    "detect_mime",
    "extract_text",
    "extract_and_ingest",
]
